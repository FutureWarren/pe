"""Document parsing for supported local source files."""

from __future__ import annotations

import csv
import re
import sys
from pathlib import Path
from typing import Iterable, Optional

from openpyxl import load_workbook
from openpyxl.utils import get_column_letter

from app.models.source import SourceDocument, SourceManifest, SourceSegment

# A bare year is 19xx / 20xx — NOT any 4-digit run (which would match data values
# like 1000 or 3500 and cause real data rows to be mistaken for period headers).
PERIOD_HEADER_RE = re.compile(
    r"(?i)(fy\s*\d{4}|\bq[1-4]\s*\d{4}\b|\b(?:19|20)\d{2}\b|\b(?:jan|feb|mar|apr|may|jun|jul|aug|sep|sept|oct|nov|dec)[a-z]*[\s\-_/]+\d{4}\b|\b\d{4}[-_/]\d{2}\b)"
)

# Bounded resource caps so a single malformed / adversarial file cannot exhaust
# memory or spin forever. When a cap is hit we still return what we parsed plus
# an explicit truncation note — never a silent drop.
MAX_PDF_PAGES = 1_000
MAX_TABLE_ROWS = 50_000


def parse_documents(manifest: SourceManifest) -> list[SourceSegment]:
    """Parse source documents into trackable segments.

    Each document is parsed in isolation: a crash on one file (encrypted PDF,
    truncated workbook, mislabeled extension, zip bomb) is converted into a
    flagged note segment so the remaining files in the data room still process.
    """

    segments: list[SourceSegment] = []
    for document in manifest.documents:
        try:
            segments.extend(_parse_document(document))
        except Exception as exc:  # noqa: BLE001 - isolate per-file parser failures
            segments.append(
                _build_note_segment(
                    document,
                    reason=f"Parser failed for {document.file_name}: {type(exc).__name__}: {exc}",
                )
            )
    return segments


def _parse_document(document: SourceDocument) -> list[SourceSegment]:
    """Parse a single source document into source-linked segments."""

    if document.file_type == "txt":
        return _parse_text_document(document)
    if document.file_type == "pdf":
        return _parse_pdf_document(document)
    if document.file_type == "csv":
        return _parse_csv_document(document)
    if document.file_type == "xlsx":
        return _parse_xlsx_document(document)
    if document.file_type == "xls":
        return _parse_xls_document(document)
    if document.file_type == "docx":
        return _parse_docx_document(document)
    return [
        _build_note_segment(
            document,
            reason=f"No parser implemented for file type {document.file_type}.",
        )
    ]


def _parse_text_document(document: SourceDocument) -> list[SourceSegment]:
    """Parse a plain text file into text sections."""

    # utf-8-sig strips a BOM if present; errors="replace" makes any undecodable
    # byte visible (U+FFFD) instead of silently deleting characters.
    text = document.absolute_path.read_text(encoding="utf-8-sig", errors="replace")
    chunks = [chunk.strip() for chunk in re.split(r"\n\s*\n", text) if chunk.strip()]
    if not chunks:
        return [_build_note_segment(document, reason="Text file was empty.")]

    segments: list[SourceSegment] = []
    for index, chunk in enumerate(chunks, start=1):
        segments.append(
            SourceSegment(
                segment_id=f"{document.source_id}-txt-{index}",
                source_id=document.source_id,
                segment_type="text_section",
                page_number=None,
                sheet_name=None,
                cell_range=None,
                section_name=f"section_{index}",
                row_number=None,
                locator_label=f"{document.file_name} section {index}",
                content=chunk,
                parsed_artifact_path=None,
                metadata={"file_name": document.file_name},
            )
        )
    return segments


def _parse_pdf_document(document: SourceDocument) -> list[SourceSegment]:
    """Parse a PDF into one text segment per page."""

    try:
        from pypdf import PdfReader
    except ImportError:
        return [_build_note_segment(document, reason="pypdf is not installed.")]

    reader = PdfReader(str(document.absolute_path))
    segments: list[SourceSegment] = []
    truncated = False
    for page_number, page in enumerate(reader.pages, start=1):
        if page_number > MAX_PDF_PAGES:
            truncated = True
            break
        text = (page.extract_text() or "").strip()
        if not text:
            continue
        segments.append(
            SourceSegment(
                segment_id=f"{document.source_id}-pdf-{page_number}",
                source_id=document.source_id,
                segment_type="page_text",
                page_number=page_number,
                sheet_name=None,
                cell_range=None,
                section_name=None,
                row_number=None,
                locator_label=f"{document.file_name} p.{page_number}",
                content=text,
                parsed_artifact_path=None,
                metadata={"file_name": document.file_name},
            )
        )

    if truncated:
        segments.append(
            _build_note_segment(
                document,
                reason=f"PDF truncated at {MAX_PDF_PAGES} pages; remaining pages were not parsed.",
            )
        )
    if not segments:
        return [_build_note_segment(document, reason="No extractable PDF text found.")]
    return segments


def _parse_csv_document(document: SourceDocument) -> list[SourceSegment]:
    """Parse CSV rows with header tracking for period-aware extraction."""

    # utf-8-sig strips a BOM (otherwise the first header cell becomes "﻿X"
    # and header detection breaks); errors="replace" keeps undecodable bytes
    # visible instead of silently deleting them.
    with document.absolute_path.open("r", encoding="utf-8-sig", errors="replace", newline="") as handle:
        sample = handle.read(65536)
        handle.seek(0)
        delimiter = _sniff_csv_delimiter(sample)
        # Widen the field-size limit (bounded) so a single oversized pasted cell
        # does not raise "_csv.Error: field larger than field limit".
        _widen_csv_field_limit()
        rows = list(csv.reader(handle, delimiter=delimiter))
    return _table_rows_to_segments(document, rows, sheet_name="CSV", segment_type="csv_row") or [
        _build_note_segment(document, reason="CSV contained no data rows.")
    ]


def _parse_xlsx_document(document: SourceDocument) -> list[SourceSegment]:
    """Parse XLSX or XLSM files into sheet row segments."""

    workbook = load_workbook(document.absolute_path, data_only=True, read_only=True)
    segments: list[SourceSegment] = []
    for sheet in workbook.worksheets:
        rows = [
            ["" if value is None else str(value) for value in row]
            for row in sheet.iter_rows(values_only=True)
        ]
        segments.extend(
            _table_rows_to_segments(
                document,
                rows,
                sheet_name=sheet.title,
                segment_type="sheet_row",
            )
        )
    if segments:
        return segments
    # No values came back. Distinguish a genuinely empty workbook from one whose
    # cells are formulas with no cached result (files written by openpyxl or some
    # LibreOffice/script exports) — the latter would otherwise drop every computed
    # number silently and look like a clean-but-empty parse.
    if _xlsx_has_formulas(document.absolute_path):
        return [
            _build_note_segment(
                document,
                reason=(
                    "Workbook cells are formulas with no cached values, so no numbers "
                    "could be read. Re-save it in Excel (which stores results) or supply "
                    "a values-only export."
                ),
            )
        ]
    return [_build_note_segment(document, reason="Workbook contained no populated rows.")]


def _parse_xls_document(document: SourceDocument) -> list[SourceSegment]:
    """Parse legacy XLS files using pandas/xlrd if available."""

    try:
        import pandas as pd
    except ImportError:
        return [_build_note_segment(document, reason="pandas is not installed.")]

    try:
        sheets = pd.read_excel(
            document.absolute_path,
            sheet_name=None,
            header=None,
            dtype=object,
        )
    except Exception as exc:  # pragma: no cover - depends on local xlrd stack
        return [_build_note_segment(document, reason=f"Unable to parse XLS file: {exc}")]

    segments: list[SourceSegment] = []
    for sheet_name, frame in sheets.items():
        rows = [
            ["" if value is None else str(value) for value in row]
            for row in frame.fillna("").values.tolist()
        ]
        segments.extend(
            _table_rows_to_segments(
                document,
                rows,
                sheet_name=sheet_name,
                segment_type="sheet_row",
            )
        )
    return segments or [_build_note_segment(document, reason="Workbook contained no populated rows.")]


def _parse_docx_document(document: SourceDocument) -> list[SourceSegment]:
    """Parse DOCX paragraphs and tables into source segments."""

    try:
        from docx import Document
    except ImportError:
        return [_build_note_segment(document, reason="python-docx is not installed.")]

    doc = Document(document.absolute_path)
    segments: list[SourceSegment] = []

    for index, paragraph in enumerate(doc.paragraphs, start=1):
        text = paragraph.text.strip()
        if not text:
            continue
        segments.append(
            SourceSegment(
                segment_id=f"{document.source_id}-docx-p-{index}",
                source_id=document.source_id,
                segment_type="docx_paragraph",
                page_number=None,
                sheet_name=None,
                cell_range=None,
                section_name=f"paragraph_{index}",
                row_number=None,
                locator_label=f"{document.file_name} paragraph {index}",
                content=text,
                parsed_artifact_path=None,
                metadata={"file_name": document.file_name},
            )
        )

    for table_index, table in enumerate(doc.tables, start=1):
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        segments.extend(
            _table_rows_to_segments(
                document,
                rows,
                sheet_name=f"table_{table_index}",
                segment_type="docx_table_row",
            )
        )

    return segments or [_build_note_segment(document, reason="DOCX contained no readable text.")]


def _table_rows_to_segments(
    document: SourceDocument,
    rows: Iterable[list[str]],
    sheet_name: str,
    segment_type: str,
) -> list[SourceSegment]:
    """Convert table-like rows into deterministic row segments."""

    segments: list[SourceSegment] = []
    current_header: Optional[list[str]] = None
    truncated = False

    for row_index, row in enumerate(rows, start=1):
        if row_index > MAX_TABLE_ROWS:
            truncated = True
            break
        cells = [cell.strip() for cell in row]
        if not any(cells):
            continue

        if _looks_like_header_row(cells):
            current_header = cells
            continue

        cell_range = _build_row_range(row_index, len(cells))
        locator_label = f"{document.file_name} {sheet_name}!{cell_range}"
        content = " | ".join(cells)
        segments.append(
            SourceSegment(
                segment_id=f"{document.source_id}-{sheet_name}-{row_index}",
                source_id=document.source_id,
                segment_type=segment_type,
                page_number=None,
                sheet_name=sheet_name,
                cell_range=cell_range,
                section_name=sheet_name,
                row_number=row_index,
                locator_label=locator_label,
                content=content,
                parsed_artifact_path=None,
                metadata={
                    "file_name": document.file_name,
                    "row_values": cells,
                    "header_values": current_header or [],
                },
            )
        )

    if truncated:
        segments.append(
            _build_note_segment(
                document,
                reason=(
                    f"{sheet_name} truncated at {MAX_TABLE_ROWS} rows; "
                    "remaining rows were not parsed."
                ),
            )
        )
    return segments


def _looks_like_header_row(cells: list[str]) -> bool:
    """Return True when a row appears to be a period header row.

    A header's period cells should be period *labels* (``FY2024``, ``Q1-24``,
    ``Jan-25``) rather than plain numbers. When every period-matching cell is a
    bare number — indistinguishable from a data value like ``2024`` used as an
    amount — only treat the row as a header if the first (label) cell is blank,
    the classic "empty corner" header layout. Otherwise it is a data row and
    must not be silently dropped.
    """

    if len(cells) < 2 or _is_numeric_like(cells[0]):
        return False
    trailing = [cell.strip() for cell in cells[1:] if cell.strip()]
    if not trailing:
        return False
    labelled = [cell for cell in trailing if PERIOD_HEADER_RE.search(cell)]
    if not labelled:
        return False
    non_bare = [cell for cell in labelled if not _is_numeric_like(cell)]
    if non_bare:
        return len(labelled) >= max(1, (len(trailing) + 1) // 2)
    # All period matches are bare numbers → header only if the corner cell is empty.
    return not cells[0].strip()


def _build_row_range(row_number: int, cell_count: int) -> str:
    """Return an Excel-like cell range for a row."""

    end_column = get_column_letter(max(cell_count, 1))
    return f"A{row_number}:{end_column}{row_number}"


def _build_note_segment(document: SourceDocument, reason: str) -> SourceSegment:
    """Build a note segment for parse failures or empty files."""

    return SourceSegment(
        segment_id=f"{document.source_id}-note",
        source_id=document.source_id,
        segment_type="file_note",
        page_number=None,
        sheet_name=None,
        cell_range=None,
        section_name=None,
        row_number=None,
        locator_label=document.file_name,
        content=reason,
        parsed_artifact_path=None,
        metadata={"file_name": document.file_name, "parse_error": reason},
    )


def _xlsx_has_formulas(path: Path, scan_limit: int = 5000) -> bool:
    """Return True if the workbook contains any formula cell (bounded scan)."""

    try:
        formula_wb = load_workbook(path, data_only=False, read_only=True)
    except Exception:  # noqa: BLE001 - detection is best-effort
        return False
    scanned = 0
    for sheet in formula_wb.worksheets:
        for row in sheet.iter_rows(values_only=True):
            for value in row:
                if isinstance(value, str) and value.startswith("="):
                    return True
                scanned += 1
                if scanned >= scan_limit:
                    return False
    return False


def _sniff_csv_delimiter(sample: str) -> str:
    """Detect the CSV delimiter, defaulting to comma.

    European exports use ``;`` (comma is their decimal separator); bank/ERP
    exports often use tab or ``|``. Without this a semicolon/tab file collapses
    into a single column and every value becomes unusable with no flag.
    """

    if not sample.strip():
        return ","
    try:
        dialect = csv.Sniffer().sniff(sample, delimiters=",;\t|")
        return dialect.delimiter
    except csv.Error:
        return ","


def _widen_csv_field_limit() -> None:
    """Raise the csv field-size limit to a bounded ceiling (not unbounded)."""

    ceiling = min(sys.maxsize, 16 * 1024 * 1024)
    try:
        if csv.field_size_limit() < ceiling:
            csv.field_size_limit(ceiling)
    except OverflowError:  # pragma: no cover - platform dependent
        csv.field_size_limit(16 * 1024 * 1024)


def _is_numeric_like(value: str) -> bool:
    """Return True when the provided string mostly looks numeric."""

    stripped = value.strip().replace(",", "").replace("$", "").replace("%", "")
    if stripped.startswith("("):
        stripped = stripped[1:]
    if stripped.endswith(")"):
        stripped = stripped[:-1]
    if not stripped:
        return False
    try:
        float(stripped)
    except ValueError:
        return False
    return True
